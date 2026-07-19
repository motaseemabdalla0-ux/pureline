"""Ingestion service: extract plain text + metadata from knowledge-base files.

Supported formats:
    .txt  - read directly (utf-8, errors ignored)
    .md   - read directly; light markdown stripping (headings/emphasis markers)
    .pdf  - extracted with pypdf
    .docx - extracted with python-docx

Each ingested file returns an ``ExtractedDocument`` carrying the plain text and
metadata (filename, file type, modified time, sha256 content hash).
"""
from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass, field

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class ExtractedDocument:
    filename: str          # base name, e.g. "irrigation-systems.md"
    file_type: str         # extension without dot, e.g. "md"
    text: str              # extracted plain text
    content_hash: str      # sha256 of the raw file bytes
    modified_time: float   # os.path.getmtime
    path: str = ""
    error: str | None = None
    meta: dict = field(default_factory=dict)


def sha256_of_file(path: str) -> str:
    """Return the sha256 hex digest of a file's raw bytes."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def _strip_markdown(text: str) -> str:
    """Lightly strip heavy markdown syntax so headings/emphasis don't pollute
    embeddings. Content words are preserved; only decorative markers removed."""
    # Remove code fences markers but keep the code text.
    text = re.sub(r"```[a-zA-Z0-9]*", "", text)
    # Headings: drop leading #'s.
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
    # Bold/italic/inline-code markers.
    text = re.sub(r"[*_`]{1,3}", "", text)
    # Links: [label](url) -> label
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    # Bullet markers.
    text = re.sub(r"^\s{0,3}[-+*]\s+", "", text, flags=re.MULTILINE)
    return text


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_md(path: str) -> str:
    return _strip_markdown(_extract_txt(path))


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page shouldn't kill the file
            continue
    return "\n".join(parts)


def _extract_docx(path: str) -> str:
    import docx  # python-docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    # Include table cell text too (FAQs are sometimes tabular).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


_EXTRACTORS = {
    ".txt": _extract_txt,
    ".md": _extract_md,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def ingest_file(path: str) -> ExtractedDocument:
    """Extract text + metadata from a single file. Never raises for parse
    errors - instead returns an ExtractedDocument with ``error`` set."""
    filename = os.path.basename(path)
    ext = os.path.splitext(filename)[1].lower()
    file_type = ext.lstrip(".")
    content_hash = sha256_of_file(path)
    modified_time = os.path.getmtime(path)

    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        return ExtractedDocument(
            filename=filename, file_type=file_type, text="",
            content_hash=content_hash, modified_time=modified_time,
            path=path, error=f"unsupported extension: {ext}",
        )
    try:
        text = extractor(path) or ""
    except Exception as exc:  # noqa: BLE001
        return ExtractedDocument(
            filename=filename, file_type=file_type, text="",
            content_hash=content_hash, modified_time=modified_time,
            path=path, error=f"{type(exc).__name__}: {exc}",
        )

    # Normalise whitespace.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return ExtractedDocument(
        filename=filename, file_type=file_type, text=text,
        content_hash=content_hash, modified_time=modified_time, path=path,
    )


def is_supported(path: str) -> bool:
    return os.path.splitext(path)[1].lower() in SUPPORTED_EXTENSIONS
